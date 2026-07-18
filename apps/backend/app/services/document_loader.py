from __future__ import annotations

import csv
import io
import json
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class UnsupportedDocumentError(ValueError):
    pass


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise UnsupportedDocumentError("Unable to decode text document")


def extract_document_text(filename: str, content_type: str | None, data: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    mime = (content_type or "").lower()

    if suffix in {".txt", ".md"} or mime.startswith("text/plain") or mime == "text/markdown":
        return _decode_text(data).strip()

    if suffix == ".json" or mime == "application/json":
        parsed = json.loads(_decode_text(data))
        return json.dumps(parsed, ensure_ascii=False, indent=2)

    if suffix == ".csv" or mime in {"text/csv", "application/csv"}:
        text = _decode_text(data)
        rows = csv.reader(io.StringIO(text))
        return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows).strip()

    if suffix == ".pdf" or mime == "application/pdf":
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()

    if (
        suffix == ".docx"
        or mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        document = DocxDocument(io.BytesIO(data))
        blocks = [
            paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n\n".join(blocks).strip()

    raise UnsupportedDocumentError(
        "Unsupported document. Allowed types: PDF, DOCX, TXT, Markdown, CSV and JSON"
    )
